import streamlit as st
from docx import Document
import docx 
import re 

EMAIL_REGEX = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'

#function for extraction of headings and text 
def extract_headings_and_text(doc):
    headings = []
    text_related_to_headings = []
    current_heading = None
    current_text = []
    email_ids = set()

    # Iterate through each paragraph in the document
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        email_matches = re.findall(EMAIL_REGEX, text)
        for email in email_matches:
            email_ids.add(email)
        # Check if the paragraph style is a heading (e.g., 'Heading 1', 'Heading 2', etc.)
        if paragraph.style.name.startswith('Heading'):
            # If we were collecting text for a previous heading, store it
            if current_heading:
                headings.append(current_heading)
                text_related_to_headings.append("\n".join(current_text))
            
            # Start a new heading
            current_heading = paragraph.text
            current_text = []
        else:
            # Collect the text under the current heading
            current_text.append(paragraph.text)

    # Add the last heading and its text
    if current_heading:
        headings.append(current_heading)
        text_related_to_headings.append("\n".join(current_text))
    
    return headings, text_related_to_headings, list(email_ids)


#function for extraction of complete text 
def extract_text_from_docx(file_path):
    """Extract text from a Word document."""
    doc = Document(file_path)
    text = []
    for para in doc.paragraphs:
        text.append(para.text)
    return '\n'.join(text)

# Define different page functions
def page_home():
    st.title("Home Page")
    #tell about this app
    st.write("Welcome to the home page of parsing  app!")
    st.write('There is two options of parsing given here....')
    st.write("Upload a document and see the magic results....")

def page_extract_word():
    st.title("Parsing of all the text from document")
    st.write("Please upload a word file and see the magic of content displayed...")
    
    # Call your extraction function here (can replace with the code I shared earlier)
    # For example, a dummy functionality:
    uploaded_file = st.file_uploader("Upload a Word Document", type="docx")
    if uploaded_file is not None:
        st.write("Extracted content will be displayed here....")
        # Extract text from the uploaded Word file
        extracted_text = extract_text_from_docx("temp.docx")

        # Display the extracted text
        st.subheader("Extracted Text:")
        st.text_area("Content", extracted_text, height=1000)
        
        
def page_extract_heading():
    st.title("Parsing of all the text and headings from document")
    st.write("Please upload a word file and see the magic of contents displayed...")
    
    # Call your extraction function here (can replace with the code I shared earlier)
    # For example, a dummy functionality:
    uploaded_file = st.file_uploader("Upload a Word Document", type="docx")
    if uploaded_file is not None:
        st.write("Extracted content will be displayed here.")
        doc = docx.Document(uploaded_file)
        
        # Extract headings and text
        headings, texts,email = extract_headings_and_text(doc)
        
        # Display the extracted content
        #st.write(hellow)
        st.write(f"**Email IDs:** {', '.join(email) if email else 'Not found'}")
        for heading, text in zip(headings, texts):
            st.subheader(heading)
            st.write(text)
            
    
# Main function that handles page navigation
def main():
    st.sidebar.title("Navigation")
    
    # Create a list of pages
    pages = {
        "Home": page_home,
        "Extract Word": page_extract_word,
        "Extract Heading":page_extract_heading
    }
    
    # Use a selectbox in the sidebar to choose a page
    selected_page = st.sidebar.selectbox("Choose a page", list(pages.keys()))
    
    # Call the selected page function
    pages[selected_page]()

if __name__ == "__main__":
    main()
